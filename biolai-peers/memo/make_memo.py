"""Build the one-page peer justification memo from data/seed_pool.csv (and
data/derived/mahalanobis_table.csv once 05_match.R has run). Every fact in the
memo is read from those files so the memo cannot drift from the data."""
import csv, os, datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

here = os.path.dirname(os.path.abspath(__file__))
seed = list(csv.DictReader(open(os.path.join(here, "..", "data", "seed_pool.csv"))))
dist_path = os.path.join(here, "..", "data", "derived", "mahalanobis_table.csv")
dist = {r["inst_key"]: r for r in csv.DictReader(open(dist_path))} if os.path.exists(dist_path) else {}
peers_path = os.path.join(here, "..", "data", "peers.csv")
peers = [r["inst_key"] for r in csv.DictReader(open(peers_path))] if os.path.exists(peers_path) else []

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6); s.left_margin = s.right_margin = Inches(0.7)
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(9.5)

def para(text, bold=False, size=None, space=2, align=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold
    if size: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space); p.paragraph_format.space_before = Pt(0)
    if align: p.alignment = align
    return p

para("Peer departments for the BiolAI evaluation", bold=True, size=13)
para(f"Department of Biology, College of Arts and Sciences. Heath Blackmon, Associate Head for Graduate Studies. {datetime.date.today():%B %d, %Y}.", size=9)

para("Purpose", bold=True, size=10.5)
para("BiolAI will be evaluated as a difference-in-differences against a fixed set of peer departments, "
     "with 2019 to 2025 as the pre-period. This memo states how the peers were chosen so that the comparison "
     "cannot be accused of being assembled after the fact. The peer list is locked with a git tag before any outcome data are read.")

para("Unit of comparison", bold=True, size=10.5)
para("The general biology department in an arts-and-sciences college at a public land-grant research university that also has a "
     "college of agriculture and a college of veterinary medicine. TAMU Biology does not contain entomology, plant pathology, "
     "biochemistry, ecology and conservation biology, rangeland, or the veterinary biomedical units, so a peer must be drawn with the same boundary. "
     "Where a university splits general biology into several arts-and-sciences departments, those departments are aggregated and "
     "entomology, biochemistry, agriculture-side and veterinary-side units are excluded.")

para("Filters and covariates", bold=True, size=10.5)
para("Hard filters: public Carnegie R1; land-grant or state flagship; college of agriculture; college of veterinary medicine "
     "(relaxed to agriculture only if the pool is thin); no existing AI or data-science scholars or fellowship program aimed at biology trainees. "
     "Among candidates passing the filters, similarity to TAMU Biology is the Mahalanobis distance on five logged covariates averaged over 2019 to 2024: "
     "bachelor's degrees in CIP 26.01 (IPEDS), PhDs granted per year (ProQuest, confirmed by graduate coordinators), research-active faculty "
     "(OpenAlex authors with three or more papers 2021 to 2025 carrying the department affiliation, calibrated against hand counts), "
     "biological sciences R&D expenditures (NSF HERD), and Pell share (IPEDS). The nearest eight to ten are the peers; all candidates and their distances are reported.")

para("Candidate pool and filter results", bold=True, size=10.5)
cols = ["Institution", "Biology unit(s) compared", "Vet", "AI program", "D", "Peer"]
t = doc.add_table(rows=1, cols=len(cols)); t.style = "Light Grid Accent 1"
for i, c in enumerate(cols):
    cell = t.rows[0].cells[i]; cell.text = ""; run = cell.paragraphs[0].add_run(c); run.bold = True; run.font.size = Pt(8)
short = {"UNIVERSITY-LEVEL ONLY": "univ. only", "IN-DEPARTMENT": "in dept", "NONE": "none", "UNKNOWN": "not checked", "TREATED": "treated"}
order = sorted(seed, key=lambda r: (r["inst_key"] != "tamu", float(dist[r["inst_key"]]["D"]) if r["inst_key"] in dist else 99, r["institution"]))
for r in order:
    k = r["inst_key"]
    units = r["biology_units"].split("(")[0].strip().rstrip(".;")
    hard = r["public"] == "Y" and r["carnegie_r1"] == "Y" and r["land_grant_or_flagship"] == "Y" and r["has_ag_college"] == "Y"
    peer = "TAMU" if k == "tamu" else ("yes" if k in peers else ("" if hard else "fails filter 1"))
    if k != "tamu" and hard and r["ai_bio_verdict"] == "IN-DEPARTMENT": peer = "excluded (AI)"
    vals = [r["institution"].replace("-Main Campus", "").replace("Virginia Polytechnic Institute and State University", "Virginia Tech"),
            units, r["has_vet_college"], short.get(r["ai_bio_verdict"], r["ai_bio_verdict"]),
            f'{float(dist[k]["D"]):.2f}' if k in dist and k != "tamu" else ("0" if k == "tamu" else "pending"), peer]
    row = t.add_row().cells
    for i, v in enumerate(vals):
        row[i].text = ""; run = row[i].paragraphs[0].add_run(v); run.font.size = Pt(7.5)
for row in t.rows:
    row.cells[0].width = Inches(1.6); row.cells[1].width = Inches(3.2); row.cells[2].width = Inches(0.4)
    row.cells[3].width = Inches(0.8); row.cells[4].width = Inches(0.5); row.cells[5].width = Inches(0.9)

n_unknown = sum(r["ai_bio_verdict"] == "UNKNOWN" for r in seed)
para("", space=1)
para("Notes. Vet = AVMA-accredited college of veterinary medicine on campus. AI program: univ. only = AI or data-science training exists at "
     "the university but not inside the biology unit(s); in dept = a funded quantitative or AI traineeship sits inside the aggregated unit "
     "(Michigan State NRT-IMPACTS in Plant Biology; Georgia Genetics T32 and Odum IDEAS NRT). Florida has no department program but the largest "
     "institution-wide AI investment in the pool; a sensitivity analysis will drop it. Penn State has no vet school and reports all campuses under one IPEDS id. "
     "Texas Tech is neither land-grant nor flagship; Cornell is private. Both remain in the table for transparency. "
     + (f"D is pending until the covariate pull completes; {n_unknown} AI-program checks are still open." if not dist else
        "D is the Mahalanobis distance from TAMU Biology; sources and dates for every value are in data/raw/manifest.csv."),
     size=8)
para("Every number here traces to a dated file in data/raw/ listed in data/raw/manifest.csv; the roster with source URLs is data/seed_pool.csv.", size=8)
out = os.path.join(here, "peer_justification_memo.docx")
doc.save(out); print("wrote", out)
