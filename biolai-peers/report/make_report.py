"""Sleuthing report: which US biology departments are most like TAMU Biology.
All numbers are embedded here with their source so the report is reproducible."""
import math, datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- data ---------------------------------------------------------------------
# PubMed affiliation counts, publication years 2021-2025, queried 2026-09-04 via the
# PubMed MCP (E-utilities). Columns: total, EEB, micro, neuro, plant (MeSH subsets).
PM = {
 "tamu":     dict(total=866, eeb=139, micro=153, neuro=100, plant=19,  aff='"Department of Biology" AND "Texas A&M University" AND "College Station"'),
 "purdue":   dict(total=918, eeb=114, micro=165, neuro=74,  plant=28,  aff='"Department of Biological Sciences" AND "Purdue University" AND "West Lafayette"'),
 "ncstate":  dict(total=838, eeb=156, micro=78,  neuro=62,  plant=18,  aff='"Department of Biological Sciences" AND "North Carolina State University"'),
 "vt":       dict(total=778, eeb=186, micro=188, neuro=74,  plant=19,  aff='"Department of Biological Sciences" AND ("Virginia Tech" OR "Virginia Polytechnic")'),
 "colostate":dict(total=774, eeb=289, micro=70,  neuro=64,  plant=90,  aff='"Department of Biology" AND "Colorado State University" AND "Fort Collins"'),
 "lsu":      dict(total=630, eeb=171, micro=93,  neuro=43,  plant=35,  aff='"Department of Biological Sciences" AND "Louisiana State University" AND "Baton Rouge"'),
 "auburn":   dict(total=495, eeb=131, micro=55,  neuro=39,  plant=15,  aff='"Department of Biological Sciences" AND "Auburn University"'),
 "wsu":      dict(total=497, eeb=147, micro=46,  neuro=28,  plant=67,  aff='"School of Biological Sciences" AND "Washington State University"'),
 "texastech":dict(total=451, eeb=125, micro=55,  neuro=27,  plant=31,  aff='"Department of Biological Sciences" AND "Texas Tech University" AND "Lubbock"'),
 "kstate":   dict(total=346, eeb=81,  micro=69,  neuro=13,  plant=25,  aff='"Division of Biology" AND "Kansas State University"'),
 "missouri": dict(total=310, eeb=43,  micro=17,  neuro=35,  plant=46,  aff='"Division of Biological Sciences" AND "University of Missouri"'),
 "florida":  dict(total=1431,eeb=487, micro=190, neuro=110, plant=127, aff='"Department of Biology" AND "University of Florida" AND "Gainesville"'),
 "pennstate":dict(total=1062, aff='"Department of Biology" AND "Pennsylvania State University" AND "University Park"'),
 "kentucky": dict(total=541,  aff='"Department of Biology" AND "University of Kentucky" AND "Lexington"'),
 "okstate":  dict(total=304,  aff='("Department of Integrative Biology" OR "Plant Biology, Ecology" OR "Department of Biology") AND "Oklahoma State University" AND "Stillwater"'),
 "msstate":  dict(total=271,  aff='"Department of Biological Sciences" AND "Mississippi State University"'),
 "oregonstate":dict(total=244, aff='"Department of Integrative Biology" AND "Oregon State University"'),
}
PM_TAMU_2016_2020 = 659

# Structural facts. bs = bachelor's in CIP 26.0101 per year (collegefactual, IPEDS-derived,
# 2021-22 unless noted); majors/phd/fac from department pages via search snippets.
D = {
 "tamu":      dict(name="Texas A&M, Department of Biology", college="Arts and Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (BS + PhD in Microbiology)", neuro="inside (BS Neuroscience; TAMIN interdisciplinary PhD)",
                   bs="742 (2020-21)", majors="2,000+ (5 majors; dept site 2025)", phd="31 PhDs 2021; ~128 grad students (Peterson's, dated)",
                   fac="~45-50 T/TT plus ~30 instructional (own knowledge; confirm)",
                   boundary="Entomology, Biochem, ECCB, Rangeland in ag college; VIBS, BIMS in vet college",
                   note="Reference. 20,000+ students served per year; largest A&P program nationally; 30% enrollment growth in 5 yr (dept strategic plan draft, Mar 2026)."),
 "purdue":    dict(name="Purdue, Department of Biological Sciences", college="Science", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (PhD research area)", neuro="inside (neurobiology area) + cross-dept institute",
                   bs="242", majors="784 majors (2016 fact sheet); 10,480 students taught 2014-15", phd="25 PhDs 2022",
                   fac="39 TT + 7 lecturers (2016 fact sheet; current likely 45-50)",
                   boundary="Entomology, Biochemistry, Botany & Plant Pathology, Forestry in College of Agriculture; Basic Medical Sciences in Vet Med",
                   note="Single broad department, molecular to ecology; PULSe umbrella PhD also feeds labs."),
 "ncstate":   dict(name="NC State, Department of Biological Sciences", college="Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (2013 merger absorbed Microbiology); plant/soil micro in CALS", neuro="no separate unit found",
                   bs="415", majors="2,100 students; 160 faculty/staff/postdocs (dept, 2013 launch)", phd="45 broad CIP 26 (year unspecified)",
                   fac="not published; est. 60-80 T/TT",
                   boundary="Entomology & Plant Pathology, Plant & Microbial Biology, Biochemistry, Applied Ecology in CALS; Molecular Biomedical Sciences in Vet Med",
                   note="Formed 2013 from Biology + Microbiology + Genetics + Env. & Molecular Toxicology; broader than TAMU by genetics and toxicology."),
 "lsu":       dict(name="LSU, Department of Biological Sciences", college="Science", ag="Y (AgCenter)", vet="Y", landgrant="Y",
                   micro="inside (BS Microbiology)", neuro="no dedicated unit found",
                   bs="334 (360 in a 2024 snippet)", majors="~2,120 majors", phd="~140 graduate students",
                   fac="59 T/TT, largest unit on campus (dept site)",
                   boundary="Entomology, Plant Pathology, Renewable Natural Resources in AgCenter/College of Agriculture; Pathobiological Sciences in Vet Med. Biochemistry is INSIDE the department",
                   note="Merger of Biochemistry, Microbiology, Plant Biology, Zoology & Physiology (c. 2011); three divisions; Museum of Natural Science."),
 "vt":        dict(name="Virginia Tech, Department of Biological Sciences", college="Science", ag="Y", vet="Y (Virginia-Maryland)", landgrant="Y",
                   micro="partly outside (Microbiology listed as separate Fralin-supported program)", neuro="outside: School of Neuroscience (College of Science)",
                   bs="250 (594 broad CIP 26)", majors="most popular major on campus", phd="27 PhDs 2022",
                   fac="not published; est. 40-50 T/TT",
                   boundary="Entomology, Biochemistry, Plant & Environmental Sciences in CALS; Fish & Wildlife in CNRE; Biomedical Sciences & Pathobiology in Vet Med",
                   note="Strong Fralin Life Sciences Institute / Roanoke biomedical pipeline; GBCB interdisciplinary PhD."),
 "auburn":    dict(name="Auburn, Department of Biological Sciences", college="Sciences and Mathematics", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (Microbial, Cellular & Molecular Biology major)", neuro="no dedicated unit found",
                   bs="315", majors="not published", phd="120+ grad students, ~65% PhD (~78)",
                   fac="not published; est. 40-50 T/TT",
                   boundary="Entomology & Plant Pathology, Fisheries in College of Agriculture; Forestry & Wildlife separate college; Chemistry & Biochemistry separate COSAM dept",
                   note="Three majors (Organismal; Microbial, Cellular & Molecular; Marine). Ichthyology/systematics strength."),
 "colostate": dict(name="Colorado State, Department of Biology", college="Natural Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="OUTSIDE: Microbiology, Immunology & Pathology in Vet Med college", neuro="OUTSIDE: Neuroscience major in Biomedical Sciences (Vet Med college)",
                   bs="318", majors="not published", phd="100+ grad students incl. MS, PSM",
                   fac="32 total, 29 full-time (Peterson's)",
                   boundary="Ag Biology (entomology, plant path) in Ag; BMB in same college; Fish, Wildlife & Conservation Biology in Warner College",
                   note="EEB-heavy (37% of papers); PhD training largely through interdepartmental Ecology and CMB programs."),
 "wsu":       dict(name="Washington State, School of Biological Sciences", college="Arts and Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="OUTSIDE: School of Molecular Biosciences (Vet Med college)", neuro="OUTSIDE: Integrative Physiology & Neuroscience (Vet Med college)",
                   bs="256", majors="not published", phd="69 grad students",
                   fac="not published",
                   boundary="Entomology, Plant Pathology, Inst. of Biological Chemistry in CAHNRS",
                   note="1999 merger of Botany + Zoology; PhDs still conferred as Botany and Zoology."),
 "texastech": dict(name="Texas Tech, Department of Biological Sciences", college="Arts and Sciences", ag="Y", vet="Y (Amarillo, accredited 2025)", landgrant="N",
                   micro="inside (BS + PhD track)", neuro="outside (Psychology; TTUHSC)",
                   bs="261", majors="not published", phd="100+ grad students",
                   fac="40-46 (TT + continuing track)",
                   boundary="Natural Resources Mgmt, Plant & Soil Science in Davis College of Ag; biochem in Chemistry & Biochemistry",
                   note="Same state, same legislature; not land-grant. Quantitative biology PhD track."),
 "kstate":    dict(name="Kansas State, Division of Biology", college="Arts and Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (own PhD in Microbiology)", neuro="not found",
                   bs="65 (26.0101 coding likely differs)", majors="~650 majors", phd="not published",
                   fac="50+ instructional and research faculty (all ranks)",
                   boundary="Entomology, Plant Pathology in College of Agriculture; Biochemistry & Molecular Biophysics in same college",
                   note="Konza Prairie LTER; studio-model intro biology. Smallest research footprint among the finalists' peers."),
 "missouri":  dict(name="Missouri, Division of Biological Sciences", college="Arts and Science", ag="Y", vet="Y", landgrant="Y",
                   micro="split (Plant, Insect & Microbial Sciences in CAFNR; Med School unit)", neuro="shared interdisciplinary program",
                   bs="210", majors="not published", phd="not published", fac="not published",
                   boundary="Biochemistry joint CAFNR/Medicine; entomology in CAFNR; Vet Pathobiology in CVM",
                   note="Bond Life Sciences Center. Small PubMed footprint (310)."),
 "florida":   dict(name="Florida, Department of Biology", college="Liberal Arts and Sciences", ag="Y (IFAS)", vet="Y", landgrant="Y",
                   micro="OUTSIDE: Microbiology & Cell Science (IFAS)", neuro="OUTSIDE: Dept of Neuroscience (Medicine), McKnight Brain Institute",
                   bs="595 (Biology BS co-administered with CALS)", majors="not published", phd="102 grad students",
                   fac="58 total faculty (grad catalog)",
                   boundary="Entomology & Nematology, Wildlife Ecology, Microbiology & Cell Science in IFAS; BMB in Medicine",
                   note="PhDs conferred as Botany and Zoology; EEB-heavy (34%); Florida Museum joint appointments; institution-wide AI initiative."),
 "okstate":   dict(name="Oklahoma State, Department of Biology (merged July 2025)", college="Arts and Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="OUTSIDE: Microbiology & Molecular Genetics (same college)", neuro="outside (OSU-CHS Tulsa)",
                   bs="100", majors="900+", phd="not published", fac="37 full-time faculty",
                   boundary="Biochem, Entomology & Plant Path, NREM in Ferguson College of Agriculture",
                   note="Brand new merger; no cell/molecular or microbiology inside."),
 "msstate":   dict(name="Mississippi State, Department of Biological Sciences", college="Arts and Sciences", ag="Y", vet="Y", landgrant="Y",
                   micro="inside (stated strength)", neuro="not found", bs="not found", majors="not published", phd="not published", fac="not published",
                   boundary="BMB + Entomology + Plant Pathology combined dept in CALS; Wildlife in Forest Resources",
                   note="Smallest PubMed footprint (271); Molecular Biology PhD sits in the ag-side biochemistry dept."),
 "oregonstate":dict(name="Oregon State, Department of Integrative Biology", college="Science (School of Life Sciences)", ag="Y", vet="Y", landgrant="Y",
                   micro="OUTSIDE: sibling Dept of Microbiology", neuro="not found", bs="191", majors="not published", phd="65 grad students",
                   fac="25 core (49 incl. affiliated)", boundary="Botany & Plant Pathology, Fisheries & Wildlife in Ag",
                   note="EEB/physiology/dev-cell only; hiring 3 TT in 2026."),
 "pennstate": dict(name="Penn State, Department of Biology", college="Eberly College of Science", ag="Y", vet="N", landgrant="Y",
                   micro="outside (BMB dept; Huck CIDD)", neuro="in-dept strength; Huck intercollege PhD", bs="253 (University Park)", majors="~1,000",
                   phd="via Huck intercollege programs", fac="not published",
                   boundary="Entomology, Plant Path & Env. Microbiology, Vet & Biomedical Sciences in College of Ag Sciences",
                   note="No vet school. PhD training is intercollege, not departmental."),
 "kentucky":  dict(name="Kentucky, Department of Biology", college="Arts and Sciences", ag="Y", vet="N", landgrant="Y",
                   micro="split (in-dept faculty; Med School dept)", neuro="split (in-dept faculty; Med School dept)", bs="173 to 211", majors="not published",
                   phd="not published", fac="not published", boundary="Entomology in College of Ag, Food & Environment",
                   note="No vet school. Urban field station."),
}

FINALISTS = ["purdue", "ncstate", "lsu", "vt", "auburn"]

def mix(k):
    p = PM[k]; t = p["total"]
    return [100 * p[x] / t for x in ("eeb", "micro", "neuro", "plant")]
def mix_dist(k):
    a, b = mix(k), mix("tamu")
    return math.sqrt(sum((x - y) ** 2 for x, y in zip(a, b)))

# ---- document -------------------------------------------------------------------
doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8); s.left_margin = s.right_margin = Inches(0.9)
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

def H(text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
def P(text, bold_lead=None, size=None, italic=False):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead + " "); r.bold = True
    r = p.add_run(text); r.italic = italic
    if size: 
        for r in p.runs: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p
def B(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead + " "); r.bold = True
    p.add_run(text); p.paragraph_format.space_after = Pt(3)
def table(header, rows, widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

title = doc.add_paragraph(); r = title.add_run("The departments most like TAMU Biology"); r.bold = True; r.font.size = Pt(18)
P(f"A sleuthing report for the BiolAI evaluation. Heath Blackmon, Department of Biology, Texas A&M University. {datetime.date.today():%B %d, %Y}.", size=9.5, italic=True)

H("Bottom line")
P("Five departments share TAMU Biology's shape: one general biology department in a science or arts-and-sciences college, "
  "at a land-grant university whose agriculture and veterinary colleges hold entomology, biochemistry, wildlife, and veterinary biomedical "
  "science, so that microbiology, neuroscience, cell biology, physiology, and ecology and evolution all live under one roof. Ranked by overall resemblance:")
B("Purdue, Department of Biological Sciences (College of Science). The closest research fingerprint in the pool and the same single-department boundary, with microbiology and neurobiology inside. Smaller undergraduate major, heavier service teaching.", "1.")
B("NC State, Department of Biological Sciences (College of Sciences). Same scale in students and papers, same land-grant sibling structure; slightly broader boundary because genetics and toxicology were folded in during a 2013 merger.", "2.")
B("LSU, Department of Biological Sciences (College of Science). The best match on identity: the largest unit on campus, 2,100 majors, 59 tenure-line faculty, microbiology inside. Broader than us because biochemistry is inside too. SEC, in the stipend-study frame.", "3.")
B("Virginia Tech, Department of Biological Sciences (College of Science). Matches on scale and sibling structure, and is the campus's most popular major; weaker on boundary because neuroscience is a separate school and microbiology is partly elsewhere.", "4.")
B("Auburn, Department of Biological Sciences (College of Sciences and Mathematics). A smaller version of us: microbiology inside, ag and vet colleges next door, SEC. Roughly 60% of our research footprint.", "5.")
P("Kansas State and Texas Tech are the alternates. Kansas State is the only other general biology unit that, like us, grants its own PhD in Microbiology, but its research footprint is 40% of ours. Texas Tech shares our state and legislature and has the same internal structure, but is not land-grant and its vet school opened in 2021.")

H("What 'like us' means")
P("TAMU Biology has seven traits that together are rare. The search looked for departments carrying as many of them as possible.")
for lead, txt in [
  ("One department, whole discipline.", "Cell and molecular, microbiology, neuroscience, physiology, ecology and evolution, and plant biology are one unit. Most large universities split these into two to five departments (UC Davis, Minnesota, Wisconsin, Illinois, Ohio State, Michigan State, Georgia, Cornell, Iowa State, Tennessee). Splitting changes governance, hiring, and teaching in ways no aggregation can undo, so the split universities were set aside at the start."),
  ("Microbiology inside, with its own degrees.", "TAMU Biology confers a BS and a PhD in Microbiology. At many land-grants, microbiology sits in the vet college (Colorado State, Washington State), the ag college (Florida), or a sibling department (Oregon State, Oklahoma State). Only Kansas State, LSU, Purdue, NC State, Auburn, Mississippi State, and Texas Tech keep it inside a general biology unit."),
  ("Neuroscience inside.", "TAMU Biology administers the Neuroscience BS and anchors the interdisciplinary PhD. Virginia Tech has a separate School of Neuroscience; Colorado State, Washington State, and Florida house it in vet or medical colleges."),
  ("Ag and vet colleges next door.", "The College of Agriculture and Life Sciences and the vet college absorb entomology, biochemistry, plant pathology, rangeland and wildlife, and veterinary biomedical science. A peer without both colleges (Penn State, Kentucky) carries some of that load in biology and is not comparable."),
  ("A very large pre-health service mission.", "2,000+ majors, 742 general-biology BS degrees a year, 20,000+ students served, the largest anatomy and physiology program nationally. This is the trait the split, research-college models (UC Davis, Minnesota) lack most."),
  ("A mid-size PhD program.", "About 31 PhDs a year and roughly 130 graduate students, with a Biology PhD and a Microbiology PhD."),
  ("A research footprint of about 170 PubMed-indexed papers a year, growing.", "866 papers carried the department affiliation in 2021 to 2025, up from 659 in 2016 to 2020. The mix is unusual: microbiology (18%) and neuroscience (12%) are large shares, ecology and evolution (16%) is smaller than at most land-grant peers, and plant biology (2%) is nearly absent because plant science sits in the ag college."),
]: B(txt, lead)

H("How the field was narrowed")
P("Twenty-five land-grant and flagship candidates were screened. Ten were removed because general biology is split across departments (listed above). Two more lack a veterinary college (Penn State, Kentucky). Of the remaining thirteen single-department candidates, seven keep microbiology inside the department, and those seven plus Colorado State, Virginia Tech, Florida, Washington State, Oregon State, and Oklahoma State were profiled on structure, scale, and research fingerprint.")
P("Two kinds of evidence drove the ranking. Structure came from department and catalog pages (college, sibling units, majors, PhD programs, mergers). Scale and research mix came from two sources that can be re-queried by anyone: PubMed affiliation counts for 2021 to 2025, with MeSH subsets for ecology and evolution, microbiology, neuroscience, and plant biology; and IPEDS-derived bachelor's counts in CIP 26.0101 (general biology) from collegefactual.com. Faculty and student headcounts come from department pages when published; several are estimates and are marked as such.")

H("Evidence table")
P("PubMed: papers listing the department in an author affiliation, publication years 2021 to 2025. Mix distance: Euclidean distance between the department's four-subfield percentage profile and TAMU's (0 = identical). BS/yr: bachelor's degrees in general biology (CIP 26.0101), 2021-22 unless noted.", size=9)
rows = []
order = ["tamu"] + FINALISTS + ["kstate", "texastech", "colostate", "florida", "wsu", "missouri", "oregonstate", "okstate", "msstate", "pennstate", "kentucky"]
for k in order:
    d = D[k]; p = PM[k]
    m = mix(k) if "eeb" in p else None
    rows.append([d["name"].replace(", Department of", ",").replace(", Division of", ",").replace(", School of", ","),
                 p["total"],
                 f"{m[0]:.0f}/{m[1]:.0f}/{m[2]:.0f}/{m[3]:.0f}" if m else "not profiled",
                 f"{mix_dist(k):.1f}" if m else "",
                 d["bs"], d["phd"], d["fac"],
                 "Y" if d["micro"].startswith("inside") else ("part" if d["micro"].startswith(("split","partly")) else "N"),
                 "Y" if d["neuro"].startswith("inside") or d["neuro"].startswith("in-dept") else ("part" if d["neuro"].startswith(("split","shared")) else "N"),
                 f'{d["ag"]}/{d["vet"]}'])
table(["Department", "PubMed 2021-25", "Mix EEB/micro/neuro/plant %", "Mix dist.", "BS/yr", "PhD program", "Faculty", "Micro in", "Neuro in", "Ag/Vet"],
      rows, widths=[1.6, 0.55, 0.9, 0.45, 0.7, 1.1, 1.1, 0.4, 0.4, 0.6], size=7.5)

H("The five, in detail")
prof = {
 "purdue": [
  "Why it is like us. One department in the College of Science spans molecular biology to ecology, and its PhD lists microbiology, neurobiology, cell, developmental, ecology and conservation, genetics, structural biology, and bioinformatics as in-house research areas. Entomology, Biochemistry, Botany and Plant Pathology, and Forestry sit in the College of Agriculture; Basic Medical Sciences and Comparative Pathobiology sit in the College of Veterinary Medicine. That is our boundary exactly. Its research fingerprint is the closest in the pool: 918 PubMed papers to our 866, with microbiology 18% (ours 18%), neuroscience 8% (ours 12%), ecology and evolution 12% (ours 16%), plants 3% (ours 2%).",
  "Where it differs. The undergraduate major is smaller (784 majors in the 2016 fact sheet; 242 general-biology BS a year against our 742), though the department taught 10,480 undergraduates in 2014-15, so the service mission is the same kind. About 25 PhDs a year to our 31. Purdue is AAU and Indiana treats Indiana University as the flagship.",
  "Verify before locking: current tenure-line count (39 in 2016) and current majors. The 2016 fact sheet PDF is on bio.purdue.edu.",
 ],
 "ncstate": [
  "Why it is like us. Created in 2013 by merging Biology, Microbiology, Genetics, and Environmental and Molecular Toxicology into one department of 2,100 students in the College of Sciences; 415 general-biology BS a year; 838 PubMed papers. The College of Agriculture and Life Sciences holds Entomology and Plant Pathology, Plant and Microbial Biology, Molecular and Structural Biochemistry, and Applied Ecology; the College of Veterinary Medicine holds Molecular Biomedical Sciences. Microbiology is inside (general and medical), with plant and soil microbiology across the street in CALS.",
  "Where it differs. Genetics and toxicology are inside, so the unit is a notch broader than ours; the Genetics and Genomics graduate program is interdepartmental. Ecology and evolution is a larger share of the papers (19%) and microbiology a smaller one (9%). NC State is not the state flagship. A 2024 NSF NRT in AI for agriculture sits in CALS plant science, not in this department, but it is the kind of exposure to watch.",
  "Verify: tenure-line faculty count (the 160 figure includes staff and postdocs) and PhDs per year in the Biology program specifically.",
 ],
 "lsu": [
  "Why it is like us. The department describes itself the way we describe ourselves: the largest academic unit on campus, 59 tenure-line faculty, about 2,120 majors, about 140 graduate students, three undergraduate majors including Microbiology, and the state's natural history museum attached. The College of Agriculture and AgCenter hold Entomology, Plant Pathology and Crop Physiology, and Renewable Natural Resources; the School of Veterinary Medicine holds Pathobiological Sciences and Comparative Biomedical Sciences. 630 PubMed papers, 334 general-biology BS a year. Flagship, land-grant, SEC, already in the 2026 stipend study frame.",
  "Where it differs. Biochemistry is inside the department (a whole division), which at TAMU is in the ag college. Neuroscience has no dedicated home. Research output per faculty member is lower than ours, and ecology and evolution is a larger share (27%).",
  "Verify: PhDs per year (not published on the site) and whether the biochemistry division should be netted out for size comparisons.",
 ],
 "vt": [
  "Why it is like us. College of Science, land-grant, ag college (Entomology, Biochemistry, Plant and Environmental Sciences) and the Virginia-Maryland vet college next door, biology the most popular major on campus, 594 broad life-science BS a year, 778 PubMed papers, 27 PhDs a year. The microbiology share of its papers (24%) is the only one in the pool higher than ours.",
  "Where it differs. Neuroscience is a separate School of Neuroscience in the same college, and microbiology graduate training is listed as a separate Fralin-supported program, so two of our in-house subfields are partly outside the department. The Fralin Biomedical Research Institute in Roanoke gives it a translational pipeline we do not have. Virginia treats UVA as the flagship.",
  "Verify: tenure-line count and how much of the 778-paper footprint comes from Fralin-affiliated faculty with joint appointments.",
 ],
 "auburn": [
  "Why it is like us. One department in the College of Sciences and Mathematics with three majors (Organismal Biology; Microbial, Cellular and Molecular Biology; Marine Biology), 315 general-biology BS a year, 120+ graduate students of whom about 78 are PhD students. Entomology and Plant Pathology and Fisheries are in the College of Agriculture, wildlife in a separate college, and the vet college is on campus. SEC, in the stipend frame, no AI or data-science trainee program found.",
  "Where it differs. About 57% of our research footprint (495 papers), with ecology and evolution at 26%. No dedicated neuroscience unit. Alabama treats the University of Alabama as flagship.",
  "Verify: tenure-line count and PhDs per year; the department publishes neither.",
 ],
}
for k in FINALISTS:
    H(D[k]["name"], 2)
    for para in prof[k]:
        lead, _, rest = para.partition(". ")
        P(rest, lead + ".")

H("Near misses, and why")
for lead, txt in [
  ("Kansas State, Division of Biology.", "The strongest structural twin at small scale: A&S college, own Microbiology PhD, entomology and plant pathology in the ag college, vet college on campus, Konza Prairie LTER. But 346 papers, about 650 majors, and 50-some faculty of all ranks make it roughly 40% of our size. Keep as the sixth peer if a sixth is wanted."),
  ("Texas Tech, Department of Biological Sciences.", "Same state, same legislature, same internal structure (microbiology inside, cell and molecular and quantitative biology tracks, 100+ graduate students, 40-46 faculty), 451 papers. Not land-grant, and the Amarillo vet school reached full accreditation only in October 2025. Politically the most useful comparison for a Texas dean; scientifically the sixth or seventh most similar."),
  ("Colorado State, Department of Biology.", "Passes every institutional filter, but microbiology and neuroscience both live in the vet college, and 37% of its papers are ecology and evolution against our 16%. It is a different kind of biology department."),
  ("Florida, Department of Biology.", "Flagship, SEC, huge (1,431 papers, 595 BS a year), but microbiology is in IFAS, neuroscience in the College of Medicine, the Biology BS is co-administered with CALS, PhDs are conferred as Botany and Zoology, and a third of its papers are ecology and evolution. Also the largest institution-wide AI investment in the country, which contaminates a BiolAI comparison."),
  ("Washington State, Oregon State, Oklahoma State.", "Each houses microbiology outside the unit (vet college, sibling department, sibling department) and is small (497, 244, 304 papers). Oklahoma State's department is two months old."),
  ("Missouri and Mississippi State.", "Correct structure, but research footprints of 310 and 271 papers, and Missouri's biochemistry and microbiology are scattered across three colleges."),
  ("Penn State and Kentucky.", "No vet school; Penn State's PhD training runs through intercollege Huck programs and its IPEDS numbers pool all campuses."),
]: B(txt, lead)

H("Confidence and what to check")
P("The rankings rest on structure (high confidence, from catalogs and department pages), PubMed counts (high confidence, exact and repeatable, but affiliation strings undercount ecology journals and any faculty who omit the department name), and headcounts (mixed: LSU, Kansas State, Oklahoma State, and Texas Tech publish theirs; Purdue's are from 2016; NC State, Virginia Tech, and Auburn do not publish tenure-line counts). Nothing here depends on a number that cannot be re-derived from the sources in the appendix.")
P("Ten minutes of work would close the main gaps: count the faculty directory pages for Purdue, NC State, Virginia Tech, and Auburn, and ask each graduate coordinator (contacts are in data/seed_pool.csv) for PhDs per year 2019 to 2025. The email draft is ready in the Gmail drafts folder.")

H("Appendix: sources")
P("PubMed queries (publication date 2021/01/01 to 2025/12/31; run 2026-09-04 through NCBI E-utilities). Subfield subsets add: EEB = (\"Biological Evolution\"[MeSH] OR \"Ecology\"[MeSH] OR \"Ecosystem\"[MeSH]); micro = (\"Bacteria\"[MeSH] OR \"Viruses\"[MeSH] OR \"Fungi\"[MeSH]); neuro = (\"Nervous System\"[MeSH] OR \"Neurons\"[MeSH] OR \"Behavior, Animal\"[MeSH]); plant = \"Plants\"[MeSH].", size=9)
table(["Department", "Affiliation query ([ad] on each phrase)", "Total"], [[D[k]["name"], PM[k]["aff"], PM[k]["total"]] for k in order], widths=[1.8, 4.2, 0.6], size=7.5)
P("Degree counts: collegefactual.com general-biology major pages per institution (IPEDS Completions, CIP 26.0101, first majors), 2021-22 unless noted. TAMU: 742 (2020-21) and 31 doctorates (2021). Department facts: department About, graduate, and catalog pages found by web search on 2026-09-04; URLs are recorded per institution in biolai-peers/data/seed_pool.csv and the research notes in biolai-peers/report/notes.md. TAMU internal figures: Department of Biology strategic-plan progress draft (March 2026) and bio.tamu.edu (2025 Class of 2029 story).", size=9)

out = __file__.replace("make_report.py", "TAMU_Biology_peer_departments.docx")
doc.save(out); print("wrote", out)
for k in order:
    if "eeb" in PM[k]: print(f"{k:12s} dist={mix_dist(k):5.1f}  mix={[round(x) for x in mix(k)]}")
